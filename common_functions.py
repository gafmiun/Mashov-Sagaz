from __future__ import annotations

import re
from typing import Any, Dict, List

import pandas as pd
from docx import Document
import docx
import docx.table
import docx.text.paragraph
from docxtpl import DocxTemplate
from bidi.algorithm import get_display

from constants import (
    MIN_GENERAL_ANSWERS,
    TOO_FEW_ANSWERS_TEXT,
    PERCENT_DECIMALS,
    DEFAULT_ZERO_VALUE,
    MIN_COMMENT_LENGTH,
    PUNCTUATION_CHARS,
    RLE,
    PDF,
    RLM,
)


def excel_to_dataframe(file_path: str) -> pd.DataFrame:
    return pd.read_excel(file_path)


def validate_excel(df: pd.DataFrame, expected_columns: List[str]) -> bool:
    actual_columns = list(df.columns)
    actual_set = set(actual_columns)
    expected_set = set(expected_columns)

    missing = expected_set - actual_set
    extra = actual_set - expected_set

    if missing:
        print("Missing required columns in Excel file.")
        print(f"Required: {sorted(expected_set)}")
        print(f"Actual:   {sorted(actual_set)}")
        print(f"Missing:  {sorted(missing)}")
        return False

    if extra:
        print("Warning: extra columns found in Excel (ignored by the code):")
        print(sorted(extra))

    if df.empty:
        print("Excel file is empty.")
        return False

    return True


def remove_parentheses_from_column(
    df: pd.DataFrame,
    column: str,
) -> pd.DataFrame:
    df[column] = df[column].astype(str).apply(
        lambda x: re.sub(r"\s*\(.*?\)", "", x).strip()
    )
    return df


def clean_dataframe(
    df: pd.DataFrame,
    commander_column: str,
) -> pd.DataFrame:
    df = df.dropna(how="all")
    df = df.dropna(subset=[commander_column])
    df = remove_parentheses_from_column(df, commander_column)
    return df


def rounded_number(x: float) -> float | int:
    if x is None:
        return DEFAULT_ZERO_VALUE
    value = float(x)
    if value.is_integer():
        return int(value)
    return round(value, PERCENT_DECIMALS)


def clean_numeric_series(series: pd.Series) -> pd.Series:
    cleaned = series.replace(r"^\s*$", pd.NA, regex=True)
    numeric = pd.to_numeric(cleaned, errors="coerce")
    return numeric.dropna()


def compute_mean_and_std(series: pd.Series) -> Dict[str, Any]:
    numeric = clean_numeric_series(series)
    n_valid = len(numeric)

    if n_valid < MIN_GENERAL_ANSWERS:
        return {
            "n": n_valid,
            "mean": TOO_FEW_ANSWERS_TEXT,
            "std": TOO_FEW_ANSWERS_TEXT,
        }

    mean_val = round(float(numeric.mean()), 2)
    std_val = numeric.std(ddof=1)
    if pd.isna(std_val):
        std_val = DEFAULT_ZERO_VALUE
    std_val = round(float(std_val), 2)

    return {"n": n_valid, "mean": mean_val, "std": std_val}


def compute_commander_general_stats(
    df_commander: pd.DataFrame,
    general_question_column: str,
) -> Dict[str, Any]:
    series = df_commander[general_question_column]
    stats = compute_mean_and_std(series)
    return {
        "average_general": stats["mean"],
        "std_general": stats["std"],
    }


def compute_mahzor_general_average(
    df_all: pd.DataFrame,
    general_question_column: str,
) -> float:
    series = clean_numeric_series(df_all[general_question_column])
    if series.empty:
        return DEFAULT_ZERO_VALUE
    return round(float(series.mean()), 2)


def rtl_embed(text: str) -> str:

    if not isinstance(text, str):
        return text

    text = text.strip()
    if not text:
        return text
    for ch in PUNCTUATION_CHARS:
        text = text.replace(ch, f"{RLM}{ch}{RLM}")

    return f"{RLE}{text}{PDF}"


def rtl_embed_graphic(text: str) -> str:

    if not isinstance(text, str):
        return text

    text = text.strip()
    if not text:
        return text

    visual = get_display(text)

    for ch in PUNCTUATION_CHARS:
        visual = visual.replace(ch, f"{RLM}{ch}{RLM}")

    return f"{RLE}{visual}{PDF}"



def replace_placeholders_in_paragraph(
    paragraph: docx.text.paragraph.Paragraph,
    values: Dict[str, Any],
) -> None:
    text_runs = [run for run in paragraph.runs if run.text]
    if not text_runs:
        return

    original_texts = [run.text for run in text_runs]
    big_text = "".join(original_texts)

    if "{{" not in big_text:
        return

    for placeholder, value in values.items():
        token = "{{" + str(placeholder) + "}}"
        if token in big_text:
            big_text = big_text.replace(token, str(value))

    pos = 0
    for run, old_text in zip(text_runs, original_texts):
        length = len(old_text)
        run.text = big_text[pos : pos + length]
        pos += length


def replace_placeholders_in_table(
    table: docx.table.Table,
    values: Dict[str, Any],
) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, values)


def replace_placeholders_in_section(
    section,
    values: Dict[str, Any],
) -> None:
    header = section.header
    footer = section.footer

    for part in (header, footer):
        if part is None:
            continue
        for paragraph in part.paragraphs:
            replace_placeholders_in_paragraph(paragraph, values)
        for table in part.tables:
            replace_placeholders_in_table(table, values)


def replace_placeholders(doc: Document, values: Dict[str, Any]) -> None:
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, values)

    for table in doc.tables:
        replace_placeholders_in_table(table, values)

    for section in doc.sections:
        replace_placeholders_in_section(section, values)


def build_basic_info_context(
    df_commander: pd.DataFrame,
    commander: str,
) -> Dict[str, Any]:
    answers_num = len(df_commander)
    return {
        "name": commander,
        "number_answers": answers_num,
    }


def build_bullet_lists_context(
    df_commander: pd.DataFrame,
    bullet_list_context: Dict[str, str],
) -> Dict[str, Dict[str, List[str]]]:
    context: Dict[str, Dict[str, List[str]]] = {}

    for key, column_name in bullet_list_context.items():
        points: List[str] = []

        if column_name in df_commander.columns:
            series = df_commander[column_name].dropna().astype(str)
            cleaned_points: List[str] = []

            for raw in series:
                text = raw.strip()
                if len(text) < MIN_COMMENT_LENGTH:
                    continue
                text = rtl_embed(text)
                cleaned_points.append(text)

            points = cleaned_points

        context[key] = {"points": points}

    return context


def merge_bullet_lists(
    df_commander: pd.DataFrame,
    commander: str,
    bullet_list_context: Dict[str, str],
) -> Dict[str, Any]:
    basic_info = build_basic_info_context(df_commander, commander)
    bullet_lists = build_bullet_lists_context(df_commander, bullet_list_context)
    return {
        **basic_info,
        **bullet_lists,
    }


def add_bullet_lists_to_template(
    tpl: DocxTemplate,
    df_commander: pd.DataFrame,
    commander: str,
    bullet_list_context: Dict[str, str],
) -> None:
    context = merge_bullet_lists(df_commander, commander, bullet_list_context)
    tpl.render(context)
