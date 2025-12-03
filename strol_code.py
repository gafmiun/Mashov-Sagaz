import pandas as pd
import docx
import re
from docx import Document
import docx.table
import docx.text.paragraph
from docx.shared import Inches, Pt, RGBColor
from typing import Dict, Union, Any
from docxtpl import DocxTemplate
import os
from typing import Optional

from constants import *

import re



def validate_excel(df: pd.DataFrame) -> bool:
    actual_columns = list(df.columns)
    actual_set = set(actual_columns)
    expected_set = set(COLUMNS)

    missing = expected_set - actual_set
    extra = actual_set - expected_set

    if missing:
        print("❌ Missing required columns in Excel file.")
        print(f"Required (expected): {sorted(expected_set)}")
        print(f"Actual:               {sorted(actual_set)}")
        print(f"Missing:              {sorted(missing)}")
        return False

    if extra:
        # Not an error – just for visibility when debugging
        print("ℹ️ Warning: extra columns found in Excel (will be ignored by the code):")
        print(f"{sorted(extra)}")

    if df.empty:
        print("❌ Excel file is empty.")
        return False

    return True




def remove_parentheses_from_commander_name(df: pd.DataFrame, column: str) -> pd.DataFrame:
    df[column] = df[column].astype(str).apply(
        lambda x: re.sub(r"\s*\(.*?\)", "", x).strip()
    )
    return df

def clean_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    df = df.dropna(how="all")
    df = df.dropna(subset=[COMMANDER_COLUMN])

    # מנקה סוגריים מכל העמודה "שם המפקד:"
    df = remove_parentheses_from_commander_name(df, COMMANDER_COLUMN)

    return df


def excel_to_dataframe(file_path=INPUT_PATH):
    return pd.read_excel(file_path)

def rounded_number(x: float):
    """
    If x is an integer return
    Otherwise return rounded float with decimals
    """
    if x is None:
        return DEFAULT_ZERO_VALUE

    if float(x).is_integer():
        return int(x)
    rounded =  round(float(x), PERCENT_DECIMALS)
    return rounded

def compute_percent(count: int, total: int) -> float:
    """
    Returns count/total * 100, rounded to PERCENT_DECIMALS .
    """
    if total <= 0:
        return DEFAULT_ZERO_VALUE
    value = (count / total) * 100
    form_value = rounded_number(value)
    return f"{form_value}%"

def compute_mahzor_general_average(df: pd.DataFrame) -> float:
    """
    Computes the overall (mahzor) average of:
        'עד כמה הייתי רוצה להיות תחת פיקודו בעתיד?'
    """
    col = GENERAL_QUESTION_COLUMN

    raw = df[col]
    series = clean_numeric_series(raw)

    if series.empty:
        return DEFAULT_ZERO_VALUE

    return round(float(series.mean()), 2)



def clean_numeric_series(series: pd.Series) -> pd.Series:
    cleaned = series.replace(r'^\s*$', pd.NA, regex=True)
    numeric = pd.to_numeric(cleaned, errors="coerce")
    return numeric.dropna()


def compute_mean_and_std(series: pd.Series) -> Dict[str, Any]:
    series = clean_numeric_series(series)
    n_valid = len(series)

    if n_valid < MIN_GENERAL_ANSWERS:
        return {
            "n": n_valid,
            "mean": TOO_FEW_ANSWERS_TEXT,
            "std": TOO_FEW_ANSWERS_TEXT,
        }

    mean_val = round(float(series.mean()), 2)
    std_val = series.std(ddof=1)
    if pd.isna(std_val):
        std_val = DEFAULT_ZERO_VALUE
    std_val = round(float(std_val), 2)

    return {"n": n_valid, "mean": mean_val, "std": std_val}

def compute_commander_general_stats(df_commander: pd.DataFrame) -> Dict[str, Any]:
    series = df_commander[GENERAL_QUESTION_COLUMN]
    stats = compute_mean_and_std(series)

    return {
        "average_general": stats["mean"],
        "std_general": stats["std"],
    }



def add_general_question_commander(df_filtered: pd.DataFrame,
                                  placeholder_to_value: Dict):
    commander_stats = compute_commander_general_stats(df_filtered)
    placeholder_to_value.update(commander_stats)


def add_general_question_mahzor(df_all: pd.DataFrame,
                               mahzor_averages: Dict):
    mahzor_avg = compute_mahzor_general_average(df_all)
    mahzor_averages["total_general"] = mahzor_avg


def count_option_in_question(
    df: pd.DataFrame,
    question_column: str,
    option_text: str,
) -> int:

    if question_column not in df.columns:
        return 0

    series = df[question_column].dropna().astype(str)
    options_for_question = QUESTION_TO_OPTIONS.get(question_column, [])
    non_none_options = [
        opt for opt in options_for_question
        if opt != NONE_OF_THE_ABOVE_OPTION
    ]

    count = 0
    for cell in series:
        text = cell.strip()
        if not text:
            continue

        if option_text == NONE_OF_THE_ABOVE_OPTION:
            # Count "none of the above" ONLY if it is the only marked option
            has_none = (NONE_OF_THE_ABOVE_OPTION in text)
            has_other = any(opt in text for opt in non_none_options)
            if has_none and not has_other:
                count += 1
        else:
            if option_text in text:
                count += 1

    return count


def get_placeholder_name_for_option(
    question_index: int,
    option_text: str,
    placeholder_type_index: int,
) -> str:

    option_key = build_option_key(option_text, question_index)
    placeholders_tuple = OPTIONS_TO_PLACEHOLDERS[option_key]
    placeholder_name = placeholders_tuple[placeholder_type_index]

    return placeholder_name


def compute_option_percentages_for_df(
    df: pd.DataFrame,
    placeholder_type_index: int,
) -> Dict[str, str]:
    results: Dict[str, str] = {}
    total_responses = len(df)

    for question_index, (question, options_for_question) in enumerate(QUESTION_TO_OPTIONS.items()):
        if question not in df.columns:
            continue
        for option_text in options_for_question:
            placeholder_name = get_placeholder_name_for_option(question_index,option_text,placeholder_type_index)
            count = count_option_in_question(
                df=df,
                question_column=question,
                option_text=option_text,
            )

            results[placeholder_name] = compute_percent(
                count,
                total_responses,
            )

    return results

def calculate_commander_multiple_choice_percentages(df_commander: pd.DataFrame, commander: str) -> Dict:
    placeholder_to_value = compute_option_percentages_for_df(
        df=df_commander,
        placeholder_type_index=PLACEHOLDER_TYPE_COMMANDER,)
    add_general_question_commander(df_commander, placeholder_to_value)

    return placeholder_to_value



def calculate_total_multiple_choice_percentage(df: pd.DataFrame) -> Dict:
    mahzor_averages = compute_option_percentages_for_df(
        df=df,
        placeholder_type_index=PLACEHOLDER_TYPE_COHORT,)
    add_general_question_mahzor(df, mahzor_averages)

    return mahzor_averages



def validate_calculations(placeholder_to_value: Dict):
    # check if there is anything left empty
    if None in placeholder_to_value.values() or "" in placeholder_to_value.values():
        print("❌ Some placeholders have empty values.")
        return False
    return True


def generate_and_fill_commander_docx(df, placeholder_to_value, commander, template_path=TEMPLATE_PATH,
                                     output_path=OUTPUT_PATH):
    # First use python docx to replace placeholders and save
    doc = Document(template_path)

    replace_placeholders(doc, placeholder_to_value)

    doc.save(output_path + commander + ".docx")

    # : Use DocxTemplate (another library) to open the processed file and add bullet lists
    add_bullet_lists(DocxTemplate(output_path + commander + ".docx"), df[df[COMMANDER_COLUMN] == commander], commander)


def replace_placeholders_in_paragraph(paragraph: docx.text.paragraph.Paragraph, values: Dict):
    text_runs = [run for run in paragraph.runs if run.text]

    if not text_runs:
        return

    original_texts = [run.text for run in text_runs]
    big_text = "".join(original_texts)

    if "{{" not in big_text:
        return

    for placeholder, value in values.items():
        token = "{{" + str(placeholder) + "}}"
        if token in big_text:
            big_text = big_text.replace(token, str(value))

    pos = 0
    for run, old_text in zip(text_runs, original_texts):
        length = len(old_text)
        run.text = big_text[pos:pos + length]
        pos += length




def replace_placeholders_in_table(table: docx.table.Table, values: Dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, values)


def replace_placeholders_in_section(section, values: Dict):
    header = section.header
    footer = section.footer
    for part in [header, footer]:
        if part is None:
            return

        for paragraph in part.paragraphs:
            replace_placeholders_in_paragraph(paragraph, values)
        for table in part.tables:
            replace_placeholders_in_table(table, values)


def replace_placeholders(doc: Document, values: Dict):
    # Body paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, values)

    # Tables
    for table in doc.tables:
        replace_placeholders_in_table(table, values)

    # Headers and footers
    for section in doc.sections:
        replace_placeholders_in_section(section, values)


def build_basic_info_context(df_commander: pd.DataFrame, commander: str) -> Dict:
    answers_num = len(df_commander)
    return {
        'name': commander,
        'number_answers': answers_num,
    }


def rtl_embed(text: str) -> str:
    "make sure text is in RTL format"
    if not isinstance(text, str):
        return text

    text = text.strip()
    if not text:
        return text


    for ch in PUNCTUATION_CHARS:
        text = text.replace(ch, f"{RLM}{ch}{RLM}")

    return f"{RLE}{text}{PDF}"


def build_bullet_lists_context(df_commander: pd.DataFrame) -> Dict:
    context: Dict[str, Dict[str, list]] = {}

    for key, column_name in BULLET_LIST_CONTEXT.items():
        points: list[str] = []

        if column_name in df_commander.columns:
            series = df_commander[column_name].dropna().astype(str)

            cleaned_points: list[str] = []
            for raw in series:
                text = raw.strip()

                #
                if len(text) < MIN_COMMENT_LENGTH:
                    continue
                text = rtl_embed(text)

                cleaned_points.append(text)

            points = cleaned_points

        context[key] = {"points": points}

    return context


def merge_bullet_lists(df_commander: pd.DataFrame, commander: str):
    basic_info = build_basic_info_context(df_commander, commander)
    bullet_lists = build_bullet_lists_context(df_commander)
    context = {
        **basic_info,
        **bullet_lists
    }
    return context
def add_bullet_lists(doc: DocxTemplate, df_commander: pd.DataFrame, commander: str):
    context = merge_bullet_lists(df_commander,commander)
    doc.render(context)
    doc.save(OUTPUT_PATH + commander + ".docx")



# ==== Excel export
def build_quantitative_header_block(
    df_commander: pd.DataFrame,
    placeholder_to_value: Dict
) -> list[list]:
    commander_name = df_commander[COMMANDER_COLUMN].iloc[0]
    num_answers = len(df_commander)

    meta_values = {
        "commander_name": commander_name,
        "num_answers": num_answers,
    }

    rows: list[list] = []
    for row_type, label, key in QUANT_HEADER_ROWS:
        if row_type == "meta":
            value = meta_values.get(key, "")
        else:
            value = placeholder_to_value.get(key, "")
        rows.append([label, value])

    return rows


def build_quantitative_table_header() -> list[str]:
    header: list[str] = [QUANT_COLUMN_QUESTION]

    for i in range(OPTIONS_PER_QUESTION):
        index = i + 1
        header.extend([
            f"{QUANT_SUBHEADER_STATEMENT} {index}",
            f"{QUANT_SUBHEADER_COMMANDER_PERCENT} {index}",
            f"{QUANT_SUBHEADER_COHORT_PERCENT} {index}",
        ])

    return header


def build_quantitative_question_row(
    question_index: int,
    question: str,
    placeholder_to_value: Dict,
    mahzor_averages: Dict
) -> list[str]:
    total_columns = 1 + OPTIONS_PER_QUESTION * OPTION_BLOCK_WIDTH
    row: list[str] = [""] * total_columns


    row[0] = question


    options_for_question = QUESTION_TO_OPTIONS.get(question, [])

    if len(options_for_question) != OPTIONS_PER_QUESTION:
        raise ValueError(
            f"Question '{question}' has {len(options_for_question)} options, "
            f"expected {OPTIONS_PER_QUESTION}."
        )

    for option_offset, option_text in enumerate(options_for_question):
        base_col = 1 + option_offset * OPTION_BLOCK_WIDTH
        row[base_col] = option_text
        if option_text == NONE_OF_THE_ABOVE_OPTION:
            option_key = f"{NONE_OF_THE_ABOVE_OPTION}_{question_index}"
        else:
            option_key = option_text

        percent_placeholder, total_placeholder = OPTIONS_TO_PLACEHOLDERS[option_key]

        commander_percent = placeholder_to_value.get(percent_placeholder, "")
        cohort_percent = mahzor_averages.get(total_placeholder, "")

        row[base_col + 1] = commander_percent
        row[base_col + 2] = cohort_percent

    return row


def build_quantitative_sheet(
    df_commander: pd.DataFrame,
    placeholder_to_value: Dict,
    mahzor_averages: Dict
) -> pd.DataFrame:
    rows: list[list] = []
    rows.extend(build_quantitative_header_block(df_commander, placeholder_to_value))
    rows.append([])
    rows.append(build_quantitative_table_header())
    total_columns = 1 + OPTIONS_PER_QUESTION * OPTION_BLOCK_WIDTH
    empty_row = [""] * total_columns

    # One row per question, with a blank row between questions
    for question_index, question in enumerate(MULTIPLE_CHOICE_COLUMNS):
        row = build_quantitative_question_row(
            question_index=question_index,
            question=question,
            placeholder_to_value=placeholder_to_value,
            mahzor_averages=mahzor_averages,
        )
        rows.append(row)
        rows.append(empty_row)

    df_quantitative = pd.DataFrame(rows)
    return df_quantitative


def collect_text_answers(df_commander: pd.DataFrame) -> Dict[str, list[str]]:
    column_to_answers: Dict[str, list[str]] = {}

    for question in OPEN_QUESTIONS_COLUMNS:
        if question in df_commander.columns:
            series = df_commander[question].dropna().astype(str)
            answers: list[str] = []

            for raw in series:
                text = raw.strip()
                if len(text) < MIN_COMMENT_LENGTH:
                    continue
                answers.append(text)
        else:
            answers = []

        column_to_answers[question] = answers

    return column_to_answers


def build_textual_sheet(df_commander: pd.DataFrame) -> pd.DataFrame:
    column_to_answers = collect_text_answers(df_commander)

    max_len = max((len(lst) for lst in column_to_answers.values()), default=0)

    data: Dict[str, list[str]] = {}

    for question, answers in column_to_answers.items():
        padded = answers + [""] * (max_len - len(answers))
        data[question] = padded

    df_textual = pd.DataFrame(data)
    return df_textual


def export_commander_excel(
    df_commander: pd.DataFrame,
    commander: str,
    placeholder_to_value: Dict,
    mahzor_averages: Dict,
    base_path: Optional[str] = None
) -> None:
    output_base = base_path if base_path is not None else COMMANDER_EXCEL_OUTPUT_PATH
    os.makedirs(output_base, exist_ok=True)

    excel_path = os.path.join(output_base, f"{commander}.xlsx")

    df_quantitative = build_quantitative_sheet(
        df_commander=df_commander,
        placeholder_to_value=placeholder_to_value,
        mahzor_averages=mahzor_averages,
    )

    df_textual = build_textual_sheet(df_commander=df_commander)

    with pd.ExcelWriter(excel_path) as writer:
        df_quantitative.to_excel(
            writer,
            sheet_name=SHEET_NAME_QUANTITATIVE,
            index=False,
            header=False,
        )
        df_textual.to_excel(
            writer,
            sheet_name=SHEET_NAME_TEXTUAL,
            index=False,
        )

    print(f"Excel for commander {commander} written to: {excel_path}")

def main(file_path=INPUT_PATH):
    try:
        df = excel_to_dataframe(file_path)
    except Exception as e:
        print(f"❌ Error reading Excel file: {e}")
        return

    if not validate_excel(df):
        print("Excel validation failed.")
        return

    df = clean_dataframe(df)

    mahzor_averages = calculate_total_multiple_choice_percentage(df)

    for commander in df[COMMANDER_COLUMN].unique():
        df_commander = df[df[COMMANDER_COLUMN] == commander]
        placeholder_to_value = calculate_commander_multiple_choice_percentages(df_commander, commander)
        # merge with mahzor averages
        placeholder_to_value.update(mahzor_averages)

        if not validate_calculations(placeholder_to_value):
            print("Calculations validation failed.")
            return

        generate_and_fill_commander_docx(df, placeholder_to_value, commander)
        export_commander_excel(
            df_commander=df_commander,
            commander=commander,
            placeholder_to_value=placeholder_to_value,
            mahzor_averages=mahzor_averages,
        )

if __name__ == "__main__":
    main(INPUT_PATH)
